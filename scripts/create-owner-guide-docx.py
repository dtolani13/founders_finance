"""Build the editable Founders Finance Owner Guide.

Install the documented dependencies first:
    py -m pip install -r scripts/requirements-owner-guide.txt

Then run:
    py scripts/create-owner-guide-docx.py
"""

from __future__ import annotations

import html
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "OWNER_GUIDE.md"
OUTPUT = ROOT / "output" / "docx" / "Founders Finance Owner Guide.docx"
RELEASE_COPY = ROOT / "release" / "Founders Finance Owner Guide.docx"
ICON = ROOT / "assets" / "brand" / "founders-finance" / "founders-finance-icon.png"

FONT = "Calibri"
CODE_FONT = "Consolas"
NAVY = "07111F"
BLUE = "0284C7"
INK = "0F172A"
SLATE = "475569"
LIGHT_BLUE = "E0F2FE"
LIGHT_SLATE = "E2E8F0"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120


def set_run_font(run, *, name=FONT, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_controls(paragraph, *, keep_with_next=False, keep_together=False, page_break_before=False):
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.paragraph_format.widow_control = True


def add_inline_runs(paragraph, text, *, size=11, color=INK):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^]]+\]\([^)]+\))")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name=CODE_FONT, size=max(size - 1, 8), color=SLATE)
        else:
            link = re.fullmatch(r"\[([^]]+)\]\(([^)]+)\)", part)
            visible = f"{link.group(1)} ({link.group(2)})" if link else part
            run = paragraph.add_run(html.unescape(visible))
            set_run_font(run, size=size, color=color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=8, color=SLATE)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_SIDE_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("end", CELL_SIDE_DXA),
    ):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    table_width = tbl_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")

    table_indent = tbl_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        tbl_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row_index, row in enumerate(table.rows):
        row_pr = row._tr.get_or_add_trPr()
        cannot_split = OxmlElement("w:cantSplit")
        row_pr.append(cannot_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_pr.append(repeat)
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def create_numbering_definition(document, *, ordered):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "\u2022")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "540")
    indentation.set(qn("w:hanging"), "270")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    run_properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    run_properties.append(color)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def create_numbering_instance(document, abstract_id):
    numbering = document.part.numbering_part.element
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    num.append(abstract_reference)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    number_properties.extend([level, number])
    paragraph_properties.append(number_properties)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True
    styles["Heading 1"].paragraph_format.page_break_before = True

    code = styles.add_style("Guide Code", 1)
    code.font.name = CODE_FONT
    code._element.rPr.rFonts.set(qn("w:ascii"), CODE_FONT)
    code._element.rPr.rFonts.set(qn("w:hAnsi"), CODE_FONT)
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor.from_string(SLATE)
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.1
    code.paragraph_format.keep_together = True
    code.paragraph_format.widow_control = True


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = paragraph.add_run("FOUNDERS FINANCE")
    set_run_font(left, size=8, color=BLUE, bold=True)
    right = paragraph.add_run("\tOWNER GUIDE")
    set_run_font(right, size=8, color=SLATE, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = paragraph.add_run("Private local-use financial operations")
    set_run_font(left, size=8, color=SLATE)
    right = paragraph.add_run("\tPage ")
    set_run_font(right, size=8, color=SLATE)
    add_page_field(paragraph)


def add_cover(document):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(32)
    if ICON.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        picture = run.add_picture(str(ICON), width=Inches(1.2))
        picture._inline.docPr.set("descr", "Founders Finance shield")
        paragraph.paragraph_format.space_after = Pt(20)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Founders Finance")
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Owner Guide")
    set_run_font(run, size=22, color=BLUE, bold=True)

    slogan = document.add_paragraph()
    slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    slogan.paragraph_format.space_after = Pt(34)
    run = slogan.add_run("WHERE CASH FLOWS. WHERE EVERY DOLLAR GOES.")
    set_run_font(run, size=10, color=SLATE, bold=True)

    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.left_indent = Inches(0.65)
    description.paragraph_format.right_indent = Inches(0.65)
    description.paragraph_format.space_after = Pt(0)
    run = description.add_run("Complete operating instructions for the private local-use installation")
    set_run_font(run, size=11, color=SLATE)
    description.add_run().add_break(WD_BREAK.PAGE)


def add_contents(document, sections):
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(18)
    run = heading.add_run("Contents")
    set_run_font(run, size=20, color=BLUE, bold=True)
    set_paragraph_controls(heading, keep_with_next=True, keep_together=True)

    for title in sections:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.12)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.1
        add_inline_runs(paragraph, title, size=10.5)
        set_paragraph_controls(paragraph, keep_together=True)


def table_widths(headers):
    lowered = [header.lower() for header in headers]
    if lowered == ["data", "location"]:
        return [3100, 6260]
    if lowered == ["export", "purpose"]:
        return [3000, 6360]
    if lowered == ["need", "go to"]:
        return [4680, 4680]
    count = len(headers)
    base = CONTENT_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(document, table_lines):
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in table_lines]
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    column_count = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (column_count - len(row)))

    table = document.add_table(rows=len(rows), cols=column_count)
    widths = table_widths(rows[0])
    set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            set_cell_fill(cell, NAVY if row_index == 0 else ("F8FAFC" if row_index % 2 == 0 else WHITE))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline_runs(paragraph, value, size=9, color=WHITE if row_index == 0 else INK)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
            set_paragraph_controls(paragraph, keep_together=True)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_markdown(document, markdown):
    lines = markdown.splitlines()
    sections = [line[3:].strip() for line in lines if line.startswith("## ")]
    add_cover(document)
    add_contents(document, sections)

    bullet_abstract = create_numbering_definition(document, ordered=False)
    ordered_abstract = create_numbering_definition(document, ordered=True)

    paragraph_lines = []
    code_lines = []
    in_code = False
    list_kind = None
    list_num_id = None
    heading_chain_remaining = 0
    index = 0

    def flush_paragraph():
        nonlocal heading_chain_remaining
        if not paragraph_lines:
            return
        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, " ".join(paragraph_lines))
        set_paragraph_controls(
            paragraph,
            keep_with_next=heading_chain_remaining > 1,
            keep_together=heading_chain_remaining > 0,
        )
        if heading_chain_remaining:
            heading_chain_remaining -= 1
        paragraph_lines.clear()

    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("```"):
            flush_paragraph()
            list_kind = None
            list_num_id = None
            if in_code:
                paragraph = document.add_paragraph(style="Guide Code")
                for code_index, code_line in enumerate(code_lines):
                    if code_index:
                        paragraph.add_run().add_break()
                    run = paragraph.add_run(code_line)
                    set_run_font(run, name=CODE_FONT, size=9, color=SLATE)
                code_lines.clear()
                in_code = False
                heading_chain_remaining = 0
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("## "):
            flush_paragraph()
            list_kind = None
            list_num_id = None
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline_runs(paragraph, line[3:].strip(), size=16, color=BLUE)
            for run in paragraph.runs:
                run.bold = True
            set_paragraph_controls(paragraph, keep_with_next=True, keep_together=True, page_break_before=True)
            heading_chain_remaining = 1
            index += 1
            continue
        if line.startswith("### "):
            flush_paragraph()
            list_kind = None
            list_num_id = None
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline_runs(paragraph, line[4:].strip(), size=13, color=BLUE)
            for run in paragraph.runs:
                run.bold = True
            set_paragraph_controls(paragraph, keep_with_next=True, keep_together=True)
            heading_chain_remaining = 2
            index += 1
            continue
        if line.startswith("#### "):
            flush_paragraph()
            list_kind = None
            list_num_id = None
            paragraph = document.add_paragraph(style="Heading 3")
            add_inline_runs(paragraph, line[5:].strip(), size=12, color=NAVY)
            for run in paragraph.runs:
                run.bold = True
            set_paragraph_controls(paragraph, keep_with_next=True, keep_together=True)
            heading_chain_remaining = 2
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|"):
            flush_paragraph()
            list_kind = None
            list_num_id = None
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            heading_chain_remaining = 0
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if bullet or numbered:
            flush_paragraph()
            kind = "ordered" if numbered else "bullet"
            if list_kind != kind:
                list_kind = kind
                list_num_id = create_numbering_instance(document, ordered_abstract if numbered else bullet_abstract)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            add_inline_runs(paragraph, (numbered or bullet).group(1))
            apply_numbering(paragraph, list_num_id)
            next_line = lines[index + 1].rstrip() if index + 1 < len(lines) else ""
            list_continues = bool(
                re.match(r"^\d+\.\s+", next_line) if numbered else re.match(r"^-\s+", next_line)
            )
            set_paragraph_controls(
                paragraph,
                keep_with_next=list_continues or heading_chain_remaining > 1,
                keep_together=True,
            )
            if heading_chain_remaining:
                heading_chain_remaining -= 1
            index += 1
            continue

        list_kind = None
        list_num_id = None
        if line.startswith(">"):
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.18)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            add_inline_runs(paragraph, line.lstrip("> "), color=SLATE)
            set_paragraph_controls(
                paragraph,
                keep_with_next=heading_chain_remaining > 1,
                keep_together=True,
            )
            if heading_chain_remaining:
                heading_chain_remaining -= 1
            index += 1
            continue
        if not line.strip():
            flush_paragraph()
            index += 1
            continue
        paragraph_lines.append(line.strip())
        index += 1

    flush_paragraph()


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    RELEASE_COPY.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_styles(document)
    configure_document(document)
    document.core_properties.title = "Founders Finance Owner Guide"
    document.core_properties.subject = "Complete owner operating instructions"
    document.core_properties.author = "Founders Finance"
    document.core_properties.keywords = "finance, owner guide, local-first, accounting operations"
    parse_markdown(document, SOURCE.read_text(encoding="utf-8"))
    document.save(OUTPUT)
    shutil.copy2(OUTPUT, RELEASE_COPY)
    print(OUTPUT)
    print(RELEASE_COPY)


if __name__ == "__main__":
    main()
